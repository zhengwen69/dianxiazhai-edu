#!/usr/bin/env python3
"""Generate 选题简介.docx from the .md source content."""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Paths ──────────────────────────────────────────────────
BASE = r"C:\Users\H1811\Desktop\CDU 固废备课文件夹2026\演示展示"
OUT_DOCX = os.path.join(BASE, "选题简介.docx")

# ─── Colors ─────────────────────────────────────────────────
DARK_BLUE = RGBColor(0x00, 0x2B, 0x5C)   # 深蓝
MID_BLUE  = RGBColor(0x00, 0x4D, 0x8C)
LIGHT_BLUE = RGBColor(0x1A, 0x6B, 0xB4)
DARK = RGBColor(0x33, 0x33, 0x33)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)      # 暗红点缀

# ─── Helpers ────────────────────────────────────────────────
def set_run_font(run, name='宋体', size=Pt(11), bold=False, color=DARK, italic=False):
    """Apply font settings to a run."""
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color

def hex_color(c):
    """Convert RGBColor to hex string."""
    return '{:02X}{:02X}{:02X}'.format(
        (c[0] if isinstance(c, tuple) else c.red) & 0xFF,
        (c[1] if isinstance(c, tuple) else c.green) & 0xFF,
        (c[2] if isinstance(c, tuple) else c.blue) & 0xFF,
    )

def add_decor_line(doc, color=DARK_BLUE, width=Inches(6.0)):
    """Add a thin decorative horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="8" w:space="1" w:color="{hex_color(color)}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def add_section_title(doc, text, level=1):
    """Add a dark blue bold section title."""
    if level == 1:
        add_decor_line(doc, DARK_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, '黑体', Pt(16), True, DARK_BLUE)
    elif level == 2:
        set_run_font(run, '黑体', Pt(13), True, DARK_BLUE)
    else:
        set_run_font(run, '黑体', Pt(11.5), True, MID_BLUE)

def add_body(doc, text, indent=False, bold_prefix=None):
    """Add body text in 宋体 11pt with 1.5 line spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # ~2 chars
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        set_run_font(run_b, '宋体', Pt(11), True, DARK)
    run = p.add_run(text)
    set_run_font(run, '宋体', Pt(11))
    return p

def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.74)
    prefix = '● ' if level == 0 else '  ○ '
    run = p.add_run(prefix + text)
    set_run_font(run, '宋体', Pt(11))

def set_cell_font(cell, text, bold=False, size=Pt(10.5), align=WD_ALIGN_PARAGRAPH.LEFT, color=DARK):
    """Set cell text with formatting."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, '宋体', size, bold, color)

def shade_cells(row, color_hex='D6E4F0'):
    """Apply background shading to all cells in a row."""
    for cell in row.cells:
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
        )
        cell._element.get_or_add_tcPr().append(shading)


# ══════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.17)
    section.right_margin  = Cm(3.17)

# Set default font for the document
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ────────────────────────────────────────────────────────────
#  COVER PAGE
# ────────────────────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()

p0 = doc.add_paragraph()
p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
r0 = p0.add_run('选题简介')
set_run_font(r0, '黑体', Pt(36), True, DARK_BLUE)
# Add a subtle underline ornament
pPr = p0._element.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:bottom w:val="single" w:sz="12" w:space="6" w:color="{hex_color(DARK_BLUE)}"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

doc.add_paragraph()

p1 = doc.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p1.add_run('《叙事·门户·智能体——')
set_run_font(r1, '华文楷体', Pt(16), True, DARK)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('　"叙·框·境·创"工程研究生教学法》')
set_run_font(r2, '华文楷体', Pt(16), True, DARK)

for _ in range(4):
    doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('黄正文　教授')
set_run_font(r3, '宋体', Pt(16), False, DARK_BLUE)

doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('成都大学环境工程系')
set_run_font(r4, '宋体', Pt(13), False, MID_BLUE)

doc.add_paragraph()

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run('二〇二六　初夏')
set_run_font(r5, '宋体', Pt(13), False, MID_BLUE)

# ── Page break ──
doc.add_page_break()

# ────────────────────────────────────────────────────────────
#  SECTION 1: 选题名称
# ────────────────────────────────────────────────────────────
add_decor_line(doc, DARK_BLUE)
add_section_title(doc, '一、选题名称')
add_body(doc, '《叙事·门户·智能体——"叙·框·境·创"工程研究生教学法》', bold_prefix='')

# ────────────────────────────────────────────────────────────
#  SECTION 2: 作者简介
# ────────────────────────────────────────────────────────────
add_section_title(doc, '二、作者简介')
add_body(doc,
    '黄正文，成都大学环境工程系教授。深耕固废治污技术与工科研究生教学凡二十余年。'
    '著环境纪实小说七部（200万字，持续更新中），连载于QQ阅读、17K平台，版权悉已登记。'
    '主成都大学研究生"课程思政"示范课《固体废弃物污染防治技术》，拟申报四川省研究生教改重点项目。')

# ────────────────────────────────────────────────────────────
#  SECTION 3: 内容提要
# ────────────────────────────────────────────────────────────
add_section_title(doc, '三、内容提要')
add_body(doc,
    '本书所论，乃一套工程研究生教学创新体系。其要有五：一曰叙事引擎，纪实小说云驱动；'
    '二曰操作模型，"叙·框·境·创"四维递进；三曰分析框架，五维错配-修复；'
    '四曰数字基座，教学-学习双重门户；五曰智能协同，九-Agent伴学体系。')
add_body(doc,
    '全书以成都大学环境工程研究生课程《固体废弃物污染防治技术》为实践载体，凡十章四附录，'
    '贯通理念建构、框架开发、门户部署、Agent协同、思政融入、八周实录、效果评价、迁移推广之全链。')

# ────────────────────────────────────────────────────────────
#  SECTION 4: 核心创新点
# ────────────────────────────────────────────────────────────
add_section_title(doc, '四、核心创新点')

innovations = [
    ('原创教学法', '：创"叙·框·境·创"四维递进教学模型——叙以启智，框以授器，境以淬能，创以致用。'),
    ('叙事-教学耦合', '：创纪实小说云驱动案例生成法。取著者七部纪实小说，以五段叙事拆解法化为标准教学案（S001-S021）。素材真切可感，版权悉已确权，自成教育资产。'),
    ('数字基座设计', '：自研"教学-学习"双重数字门户。纯静态架构（HTML/CSS/JS），GitHub Pages零费部署。零配置即得进度追踪，四段式页型贯通学习全周。'),
    ('AI Agent协同', '：设计九-Agent协同教学体系，立师端至生端双向映射之法。Agent自教具化为伴学者，不替师，唯辅学。'),
    ('课程思政叙事化', '：立"看不见的思政"之则。化思政元素于叙事素材、错配框架、场景模拟之中，浑然不觉而润物无声，破工程课堂"两张皮"困局。'),
]
for i, (prefix, text) in enumerate(innovations, 1):
    p = add_body(doc, text, indent=True, bold_prefix=f'{i}. {prefix}')
    p.paragraph_format.left_indent = Cm(0.74)

# ────────────────────────────────────────────────────────────
#  SECTION 5: 读者对象
# ────────────────────────────────────────────────────────────
add_section_title(doc, '五、读者对象')
readers = [
    '环境工程及相关工科专业高校教师、研究生导师',
    '高校教学改革与课程建设管理人员',
    '教育技术、教育信息化领域研究者与从业者',
    '关注工程教育创新与课程思政建设之研究生与本科生',
]
for r in readers:
    add_bullet(doc, r)

# ────────────────────────────────────────────────────────────
#  SECTION 6: 市场分析
# ────────────────────────────────────────────────────────────
add_section_title(doc, '六、市场分析')
add_body(doc,
    '工程教育自"知识传授型"向"能力培养型"转轨，呼声日盛。然成体系、可操作、有实据支撑之教学法专著，'
    '仍属稀缺。本书以一门课为完整样本，供理念、实施、评价全套可迁移方案。'
    '著者已出版纪实小说七部（QQ阅读、17K），为本书叙事引擎筑独特内容壁垒，兼蓄跨圈传播潜力。')

# ────────────────────────────────────────────────────────────
#  SECTION 7: 预估规模
# ────────────────────────────────────────────────────────────
add_section_title(doc, '七、预估规模')
add_body(doc, '正文约二十万字，含附录，成书约四百五十页。')

# ────────────────────────────────────────────────────────────
#  SECTION 8: 已有基础
# ────────────────────────────────────────────────────────────
add_section_title(doc, '八、已有基础')
foundations = [
    '课程已获批成都大学研究生"课程思政"示范课立项',
    '配套数字门户上线运营',
    '纪实小说七部出版上线，版权悉已登记',
    '教学素材母库建成（S001-S021，凡二十一选段）',
    '本科《固体废物处理与处置》六章知识体系已整理归档，与研究生前置自检系统联动',
    '衍生教研论文七篇，大纲齐备，部分已入定稿',
    '四川省研究生教改重点项目申报推进中',
]
for f in foundations:
    add_bullet(doc, f)

# ────────────────────────────────────────────────────────────
#  SECTION 9: 联系方式
# ────────────────────────────────────────────────────────────
add_section_title(doc, '九、联系方式')

# Contact table
contact_table = doc.add_table(rows=4, cols=2, style='Table Grid')
contact_table.alignment = WD_TABLE_ALIGNMENT.LEFT

# Set column widths
for row in contact_table.rows:
    row.cells[0].width = Cm(3.0)
    row.cells[1].width = Cm(9.5)

contact_data = [
    ('姓　名', '黄正文'),
    ('单　位', '成都大学建筑与土木工程学院环境工程系'),
    ('地　址', '成都 610106'),
    ('邮　箱', 'huangzhengwen@cdu.edu.cn'),
]
for i, (label, value) in enumerate(contact_data):
    row = contact_table.rows[i]
    shade_cells(row, 'EAF0F8')
    set_cell_font(row.cells[0], label, bold=True, size=Pt(11), align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_font(row.cells[1], value, bold=False, size=Pt(11))

# ── Author info table (at the end) ──
doc.add_paragraph()  # spacer

# ────────────────────────────────────────────────────────────
#  APPENDIX: 本书三级目录
# ────────────────────────────────────────────────────────────
doc.add_page_break()
add_decor_line(doc, DARK_BLUE)
add_section_title(doc, '附件：本书三级目录')
add_decor_line(doc, DARK_BLUE, width=Inches(3.0))

# ── Chapter 1 ──
add_section_title(doc, '第一章　绪论：工程研究生教学的困境与突围（~1.5万字）', level=2)
add_section_title(doc, '1.1　工程研究生教育的时代困境', level=3)
add_bullet(doc, '1.1.1　"三多三少"的课堂现实', level=0)
add_bullet(doc, '1.1.2　传统讲授法在工科研究生层次的三个极限', level=0)
add_bullet(doc, '1.1.3　从"技术执行者"到"系统决策者"的培养目标跃迁', level=0)
add_section_title(doc, '1.2　固体废弃物课程的教改契机', level=3)
add_bullet(doc, '1.2.1　固废课程的跨学科本质与教学挑战', level=0)
add_bullet(doc, '1.2.2　成都大学环境工程系研究生课程改革背景', level=0)
add_bullet(doc, '1.2.3　课程思政示范课程立项与国家教改政策窗口', level=0)
add_section_title(doc, '1.3　本书的写作逻辑与阅读指引', level=3)
add_bullet(doc, '1.3.1　"理念→框架→方法→平台→实践→评价→推广"的逻辑链', level=0)
add_bullet(doc, '1.3.2　理论章与实践章的差异化阅读建议', level=0)
add_bullet(doc, '1.3.3　与演示门户的配合使用说明', level=0)

# ── Chapter 2 ──
add_section_title(doc, '第二章　理论基础："叙·框·境·创"四维教学法（~2.5万字）', level=2)
add_section_title(doc, '2.1　四维教学法的理论溯源', level=3)
add_bullet(doc, '2.1.1　叙事学习理论：故事作为认知入口', level=0)
add_bullet(doc, '2.1.2　建构主义学习理论：框架作为思维工具', level=0)
add_bullet(doc, '2.1.3　情境学习理论：场景作为能力熔炉', level=0)
add_bullet(doc, '2.1.4　成果导向教育（OBE）：创造作为学习闭环', level=0)
add_section_title(doc, '2.2　"叙·框·境·创"的完整学理推导', level=3)
add_bullet(doc, '2.2.1　叙：以纪实小说触发感性认知的叙事设计', level=0)
add_bullet(doc, '2.2.2　框：以"错配-修复"框架赋予分析工具', level=0)
add_bullet(doc, '2.2.3　境：以场景模拟淬炼应用能力的沉浸式设计', level=0)
add_bullet(doc, '2.2.4　创：以三轨综合产出实现创造致用', level=0)
add_bullet(doc, '2.2.5　四维递进的内在逻辑：从认知激活到知识再生产的闭环', level=0)
add_section_title(doc, '2.3　与已有教学法的比较分析', level=3)
add_bullet(doc, '2.3.1　与案例教学法的异同', level=0)
add_bullet(doc, '2.3.2　与PBL（问题导向学习）的异同', level=0)
add_bullet(doc, '2.3.3　与翻转课堂的异同', level=0)
add_bullet(doc, '2.3.4　四维教学法的独特贡献：叙事驱动×框架穿透×场景加速×创造闭环', level=0)

# ── Chapter 3 ──
add_section_title(doc, '第三章　分析框架：五维错配-修复体系（~2.5万字）', level=2)
add_section_title(doc, '3.1　为什么是"错配"——工程实践的认识论提炼', level=3)
add_bullet(doc, '3.1.1　资源错配导致效率损失', level=0)
add_bullet(doc, '3.1.2　空间错配导致环境不公', level=0)
add_bullet(doc, '3.1.3　从"发现问题"到"定义概念"：错配作为统一分析范式', level=0)
add_section_title(doc, '3.2　七个子框架的完整展开', level=3)
add_bullet(doc, '3.2.1　空间错配五维：选址谬误 / 规划滞后 / 容量失衡 / 运输悖论 / 边界模糊', level=0)
add_bullet(doc, '3.2.2　选址正义性四维：地质承载 / 环境容量 / 社会可接受 / 代际公平', level=0)
add_bullet(doc, '3.2.3　技术错配三型：过度设计 / 低配高用 / 方向性错误', level=0)
add_bullet(doc, '3.2.4　垃圾时间三维：工艺时间 / 环境效益兑现时间 / 投资回收时间', level=0)
add_bullet(doc, '3.2.5　行为错配四因：认知阻隔 / 便利性陷阱 / 社会规范缺位 / 反馈延迟', level=0)
add_bullet(doc, '3.2.6　EPR错配三源：责任界定模糊 / 执行成本转嫁 / 监督失灵', level=0)
add_bullet(doc, '3.2.7　人文错配三维：分配正义缺陷 / 程序正义缺失 / 承认正义忽视', level=0)
add_section_title(doc, '3.3　五维统一诊断矩阵', level=3)
add_bullet(doc, '3.3.1　矩阵的构建逻辑：空间×技术×行为×人文×时间', level=0)
add_bullet(doc, '3.3.2　从独立诊断到系统整合的推演过程', level=0)
add_bullet(doc, '3.3.3　矩阵在综合案例中的应用示例', level=0)
add_section_title(doc, '3.4　"五垃圾"标签体系——课程的认知脚手架', level=3)
add_bullet(doc, '3.4.1　垃圾资产 / 垃圾空间 / 垃圾时间 / 垃圾技术 / 垃圾模式', level=0)
add_bullet(doc, '3.4.2　标签与错配维度的映射关系', level=0)
add_bullet(doc, '3.4.3　标签在八周教学中的渐进式展开', level=0)

# ── Chapter 4 ──
add_section_title(doc, '第四章　叙事引擎：纪实小说云驱动的案例方法（~2.5万字）', level=2)
add_section_title(doc, '4.1　纪实小说云的构建', level=3)
add_bullet(doc, '4.1.1　七部纪实小说的创作背景与教育意图', level=0)
add_bullet(doc, '4.1.2　S001-S021选段的选择逻辑与维度覆盖', level=0)
add_bullet(doc, '4.1.3　素材母库的维护机制与扩展规则', level=0)
add_section_title(doc, '4.2　五段叙事拆解法', level=3)
add_bullet(doc, '4.2.1　背景→冲突→决策→结局→启示的标准化拆解流程', level=0)
add_bullet(doc, '4.2.2　从"工程叙事"到"教学案例"的转化三步骤', level=0)
add_bullet(doc, '4.2.3　方法论语块：可复用的拆解模板与操作范例', level=0)
add_section_title(doc, '4.3　四重错配的叙事化呈现', level=3)
add_bullet(doc, '4.3.1　空间错配叙事：《龙栖湾》的土地贬值与邻避困境', level=0)
add_bullet(doc, '4.3.2　技术错配叙事：《青龙湖》的技术与来料双重错配', level=0)
add_bullet(doc, '4.3.3　行为错配叙事：《杨柳坝与刘家湾》的乡村分类困局', level=0)
add_bullet(doc, '4.3.4　人文错配叙事：《三多里巷》的尊严与正义', level=0)
add_bullet(doc, '4.3.5　四重叙事的方法论闭环：从个案到原型', level=0)
add_section_title(doc, '4.4　纪实小说作为教学素材的独特价值', level=3)
add_bullet(doc, '4.4.1　真实性与情感性的双重教育优势', level=0)
add_bullet(doc, '4.4.2　版权确权的教育资产属性', level=0)
add_bullet(doc, '4.4.3　与商业案例、虚构案例的比较优势', level=0)

# ── Chapter 5 ──
add_section_title(doc, '第五章　数字基座："教学-学习"双重门户设计与操作（~2.5万字）', level=2)
add_section_title(doc, '5.1　门户的设计哲学', level=3)
add_bullet(doc, '5.1.1　从"课程网站"到"教学数字门户"的理念演进', level=0)
add_bullet(doc, '5.1.2　双重门户架构：教师端"管"×学生端"学"的协同设计', level=0)
add_bullet(doc, '5.1.3　零配置部署哲学：纯静态HTML/CSS/JS + GitHub Pages推送即用', level=0)
add_bullet(doc, '5.1.4　响应式设计：移动端与桌面端的一体化体验', level=0)
add_section_title(doc, '5.2　教师门户——功能架构与使用指南', level=3)
add_bullet(doc, '5.2.1　首页导航（index.html）：信息枢纽与品牌展示', level=0)
add_bullet(doc, '5.2.2　课程概览：OBE四目标×八周框架的矩阵呈现', level=0)
add_bullet(doc, '5.2.3　章节导航：八周完整教学方案的层级索引', level=0)
add_bullet(doc, '5.2.4　知识导图：故事云驱动的知识树可视化', level=0)
add_bullet(doc, '5.2.5　Agent团队演示：9-Agent协同体系的可视化说明', level=0)
add_bullet(doc, '5.2.6　互动社区：签到打卡·周报提交·同行评价的功能设计', level=0)
add_bullet(doc, '5.2.7　教师仪表盘：数据查阅·答疑管理·教学监控', level=0)
add_section_title(doc, '5.3　学习园地——以学为中心的子门户', level=3)
add_bullet(doc, '5.3.1　信息架构设计：三折叠导览（概览/导航/自测）', level=0)
add_bullet(doc, '5.3.2　八周学习计划与进度追踪：localStorage零配置实现详解', level=0)
add_bullet(doc, '5.3.3　每周学习页的四段式纵向结构：预习→学习→作业→AI赋能', level=0)
add_bullet(doc, '5.3.4　学生操作指南：从第1周到第8周的完整学习路径', level=0)
add_bullet(doc, '5.3.5　进度标记与学习数据的自主管理', level=0)
add_bullet(doc, '5.3.6　三层渐进式前置自检导航——"温故而知新"按钮、隐藏知识点列表与锚点详情页的交互设计', level=0)
add_bullet(doc, '5.3.7　本科知识归档体系——六章知识要点HTML化、跨章锚点跳转与研-本课程联动机制', level=0)
add_section_title(doc, '5.4　技术实现与部署维护', level=3)
add_bullet(doc, '5.4.1　技术栈说明（HTML5+CSS3+原生JavaScript+localStorage）', level=0)
add_bullet(doc, '5.4.2　GitHub Pages部署与版本管理流程', level=0)
add_bullet(doc, '5.4.3　内容更新操作手册（含截图分步说明）', level=0)
add_section_title(doc, '5.5　门户功能索引与截图对照表', level=3)

# ── Chapter 6 ──
add_section_title(doc, '第六章　智能协同：AI Agent教学体系（~2万字）', level=2)
add_section_title(doc, '6.1　从单点AI工具到Agent协同体系的演进', level=3)
add_bullet(doc, '6.1.1　现有AI教育工具的碎片化困境', level=0)
add_bullet(doc, '6.1.2　Agent体系的设计初衷：全流程覆盖而非单点替代', level=0)
add_bullet(doc, '6.1.3　核心原则：Agent辅助教师，而非取代教师', level=0)
add_section_title(doc, '6.2　9-Agent协同教学体系全景', level=3)
add_bullet(doc, '6.2.1　教师端Agent的角色分工与工作流', level=0)
add_bullet(doc, '6.2.2　各Agent的能力边界与协作数据流', level=0)
add_bullet(doc, '6.2.3　Agent体系的可视化架构', level=0)
add_section_title(doc, '6.3　"教→伴学"的双向角色映射方法论', level=3)
add_bullet(doc, '6.3.1　角色再定义原则：保持编号与数量，重塑学生端职能', level=0)
add_bullet(doc, '6.3.2　教师端→学生端的逐对映射表与转换逻辑', level=0)
add_bullet(doc, '6.3.3　案例一：Agent-S——从"故事拆解员"到"案例解读员"', level=0)
add_bullet(doc, '6.3.4　案例二：Agent-B——从"学术品牌建设"到"研创向导"', level=0)
add_bullet(doc, '6.3.5　角色映射的设计约束：信息对等 / 术语一致 / 独立可调用', level=0)
add_section_title(doc, '6.4　学生端Agent的使用场景与方法', level=3)
add_bullet(doc, '6.4.1　学习流6步与Agent的对应调用关系', level=0)
add_bullet(doc, '6.4.2　专属Prompt的设计原理与使用指南', level=0)
add_bullet(doc, '6.4.3　助学团队学生版门户操作说明', level=0)

# ── Chapter 7 ──
add_section_title(doc, '第七章　价值融入：课程思政的叙事化路径（~2万字）', level=2)
add_section_title(doc, '7.1　工程课程思政的两个常见困境', level=3)
add_bullet(doc, '7.1.1　"标签化"：罗列思政要求但不落地', level=0)
add_bullet(doc, '7.1.2　"两张皮"：思政内容与专业教学分离', level=0)
add_bullet(doc, '7.1.3　叙事化路径的破解逻辑：故事比说教更能传递价值观', level=0)
add_section_title(doc, '7.2　课程思政的叙事化路径设计', level=3)
add_bullet(doc, '7.2.1　门户首页思政标语的叙事承载："绿水青山=金山银山"', level=0)
add_bullet(doc, '7.2.2　纪实小说云作为思政素材库：S001-S021的思政内涵地图', level=0)
add_bullet(doc, '7.2.3　"看不见的思政"原则：价值观嵌入故事，而非口号嵌入课堂', level=0)
add_section_title(doc, '7.3　思政元素在错配框架中的多维嵌入', level=3)
add_bullet(doc, '7.3.1　空间维度→社会公平："谁的资产被剥夺？"', level=0)
add_bullet(doc, '7.3.2　技术维度→工程师责任："技术为谁服务？"', level=0)
add_bullet(doc, '7.3.3　行为维度→制度伦理："制度设计中的公平"', level=0)
add_bullet(doc, '7.3.4　人文维度→环境正义："邻避不是自私"', level=0)
add_section_title(doc, '7.4　场景沉浸中的伦理觉醒', level=3)
add_bullet(doc, '7.4.1　模拟法庭中的价值判断与角色代入', level=0)
add_bullet(doc, '7.4.2　公众听证会中的公平诉求', level=0)
add_bullet(doc, '7.4.3　运营危机应对中的责任担当', level=0)

# ── Chapter 8 ──
add_section_title(doc, '第八章　八周实践：教学全程实录（~3万字）', level=2)
# Table for 8 weeks
toc_table = doc.add_table(rows=9, cols=4, style='Table Grid')
toc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
headers = ['节', '周次', '主题', '主导维度']
week_data = [
    ('8.1', '第1周', '空间错配Ⅰ——"叙"的启动', '叙'),
    ('8.2', '第2周', '空间错配Ⅱ——"框"的建立', '框'),
    ('8.3', '第3周', '技术错配Ⅰ——"境"的开启', '境'),
    ('8.4', '第4周', '技术错配Ⅱ——"境"的深化', '境→创'),
    ('8.5', '第5周', '行为错配Ⅰ——"境"的拓展', '境'),
    ('8.6', '第6周', '行为错配Ⅱ——框架与场景的叠加', '框→境'),
    ('8.7', '第7周', '人文错配——"境"的升华', '境'),
    ('8.8', '第8周', '思维整合——"创"的收束', '创'),
]
for j, h in enumerate(headers):
    set_cell_font(toc_table.rows[0].cells[j], h, True, Pt(10.5), WD_ALIGN_PARAGRAPH.CENTER, WHITE)
shade_cells(toc_table.rows[0], hex_color(DARK_BLUE))
for i, (sec, wk, topic, dim) in enumerate(week_data, 1):
    set_cell_font(toc_table.rows[i].cells[0], sec, False, Pt(10.5), WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_font(toc_table.rows[i].cells[1], wk, False, Pt(10.5), WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_font(toc_table.rows[i].cells[2], topic, False, Pt(10.5))
    set_cell_font(toc_table.rows[i].cells[3], dim, False, Pt(10.5), WD_ALIGN_PARAGRAPH.CENTER)
    if i % 2 == 0:
        shade_cells(toc_table.rows[i], 'F0F4FA')

doc.add_paragraph()
add_body(doc, '（每节含：教学目标与框架导入 · 课堂实录 · 教学反思）', indent=True)

# ── Chapter 9 ──
add_section_title(doc, '第九章　效果评价：数据、反思与迭代（~1.5万字）', level=2)
add_section_title(doc, '9.1　评价体系设计', level=3)
add_bullet(doc, '9.1.1　过程性评价：学习行为数据的追踪与解读', level=0)
add_bullet(doc, '9.1.2　终结性评价：开卷案例分析 + 思维重塑报告', level=0)
add_bullet(doc, '9.1.3　多维评价：同行评价 × 自我评价 × 教师评价的三角验证', level=0)
add_section_title(doc, '9.2　教学效果实证分析', level=3)
add_bullet(doc, '9.2.1　学习行为数据呈现', level=0)
add_bullet(doc, '9.2.2　学生三轨产出（论文/设计/教改）的选题分布', level=0)
add_bullet(doc, '9.2.3　与纯讲授法班级的对比分析', level=0)
add_section_title(doc, '9.3　教学反思与框架迭代', level=3)
add_bullet(doc, '9.3.1　四维教学法的优势与适用边界', level=0)
add_bullet(doc, '9.3.2　教学中的失败案例与调整过程', level=0)
add_bullet(doc, '9.3.3　框架从v1.0到当前版本的迭代日志', level=0)

# ── Chapter 10 ──
add_section_title(doc, '第十章　推广与展望：向其他工程课程的迁移（~1万字）', level=2)
add_section_title(doc, '10.1　四维教学法的普适性论证', level=3)
add_bullet(doc, '10.1.1　迁移到其他工程课程的三个前提条件', level=0)
add_bullet(doc, '10.1.2　可迁移的教学要素清单', level=0)
add_bullet(doc, '10.1.3　不可迁移的独特性边界', level=0)
add_section_title(doc, '10.2　数字门户的可迁移框架', level=3)
add_bullet(doc, '10.2.1　纯静态架构的通用性与模板化方案', level=0)
add_bullet(doc, '10.2.2　部署与维护的标准化指南', level=0)
add_section_title(doc, '10.3　纪实小说作为教育资产的前景', level=3)
add_bullet(doc, '10.3.1　小说云模式的可持续创作机制', level=0)
add_bullet(doc, '10.3.2　向其他学科扩展的可能性', level=0)
add_bullet(doc, '10.3.3　教育叙事出版：一个值得关注的蓝海', level=0)
add_section_title(doc, '10.4　结语：教学即研究，课堂即田野', level=3)

# ── Appendices ──
add_section_title(doc, '附录', level=2)
add_bullet(doc, '附录A　场景沉浸脚本全集（6个脚本）')
add_bullet(doc, '附录B　教学素材母库索引 S001-S021')
add_bullet(doc, '附录C　中英术语对照表')
add_bullet(doc, '附录D　门户截图目录 SC01-SC20')

# ── Footer note ──
doc.add_paragraph()
add_decor_line(doc, DARK_BLUE)
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_end = p_end.add_run('全稿凡二十万言')
set_run_font(r_end, '楷体', Pt(12), True, DARK_BLUE)

doc.add_paragraph()

p_sig = doc.add_paragraph()
p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_sig = p_sig.add_run('黄正文 | 成都大学环境工程系 | 二〇二六 初夏')
set_run_font(r_sig, '宋体', Pt(10), False, MID_BLUE)

# ── Save ──
doc.save(OUT_DOCX)
print(f'DOCX saved to: {OUT_DOCX}')
print(f'File size: {os.path.getsize(OUT_DOCX):,} bytes')
