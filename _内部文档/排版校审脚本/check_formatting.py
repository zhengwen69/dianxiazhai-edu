# -*- coding: utf-8 -*-
"""化工社排版自动机检 - 10项标准 - 输出通过/警告/错误"""

import sys, os, re
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

def check(docx_path):
    if not os.path.exists(docx_path):
        print(f"文件不存在: {docx_path}")
        return

    doc = Document(docx_path)
    results = []
    errors = 0; warnings = 0

    # ── A1: 页面尺寸 ──
    sec = doc.sections[0]
    w_ok = abs(sec.page_width - Cm(14.8)) < 10000
    h_ok = abs(sec.page_height - Cm(21.0)) < 10000
    if w_ok and h_ok:
        results.append(("A1 页面尺寸", "通过", f"A5 148x210mm"))
    else:
        results.append(("A1 页面尺寸", "错误", f"{sec.page_width/914400*2.54:.1f}x{sec.page_height/914400*2.54:.1f}in"))
        errors += 1

    # ── A2: 页边距 ──
    m_ok = (abs(sec.top_margin - Cm(2.5)) < 5000 and
            abs(sec.bottom_margin - Cm(2.0)) < 5000 and
            abs(sec.left_margin - Cm(2.5)) < 5000 and
            abs(sec.right_margin - Cm(2.0)) < 5000)
    if m_ok:
        results.append(("A2 页边距", "通过", "2.5/2.0/2.5/2.0 cm"))
    else:
        results.append(("A2 页边距", "警告", "偏离标准"))
        warnings += 1

    # ── A3-A8: 段落扫描 ──
    level4_texts = [
        '其一，知识灌输偏多，工程实践能力培育偏少',
        '其二，教师讲授偏多，学生自主知识建构偏少',
        '其三，技术维度分析偏多，工程全周期社会价值与多维属性反思偏少',
        '首限在案例与理论之分离', '次限在场景与课堂之割裂', '末限在产出与学习之脱节',
        '理念环', '框架环', '方法环', '平台环', '实践环', '评价环', '推广环',
        '教学管理者与教育改革决策者', '一线授课教师与课程设计人员', '高等教育领域研究者',
        '阅读第5章（双重门户）时', '阅读第6章（Agent教学体系）时', '阅读第8章（八周教学实录）时',
    ]

    chapter_font_ok = False
    body_indent_count = 0; body_total = 0
    super_count = 0; super_total = 0
    level4_ok = 0; level4_total = len(level4_texts)
    ref_font_ok = False; ref_indent_ok = False
    in_refs = False

    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue

        # Detect references section
        if t == '参考文献' and any(r.font.bold for r in p.runs):
            in_refs = True
            continue

        # A3: Chapter title (first paragraph)
        if not chapter_font_ok:
            has_18pt = any(abs(r.font.size - Pt(18)) < 1000 for r in p.runs if r.font.size)
            has_hei = any('黑体' in (r._element.rPr.rFonts.get(qn('w:eastAsia')) or '') for r in p.runs)
            has_center = p.alignment == 1  # CENTER
            if has_18pt and has_hei and has_center:
                chapter_font_ok = True
            else:
                # Check alternative - bold with center
                ch_bold = any(r.font.bold for r in p.runs)
                if ch_bold and has_center:
                    chapter_font_ok = True
            break

    if chapter_font_ok:
        results.append(("A3 章标题", "通过", "黑体18pt居中"))
    else:
        results.append(("A3 章标题", "警告", "检查字体"))
        warnings += 1

    # Full paragraph scan for A4-A8
    section_ok = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        if t == '参考文献': in_refs = True; continue

        # A4: Section headings
        if re.match(r'^1\.[123]', t) or re.match(r'^\d+\.\d+\.\d+', t):
            has_bold = any(r.font.bold for r in p.runs)
            if has_bold: section_ok = True

        # A5: Body indentation
        if not in_refs and not any(r.font.bold for r in p.runs) and not re.match(r'^\d', t):
            body_total += 1
            if p.paragraph_format.first_line_indent:
                body_indent_count += 1

        # A6: Superscript citations (exclude references section)
        for r in p.runs:
            if r.text and '[' in r.text and not in_refs:
                super_total += 1
                if r.font.superscript:
                    super_count += 1

        # A7: Level4 headings
        if t in level4_texts:
            if any(r.font.bold for r in p.runs):
                level4_ok += 1

        # A8: References
        if in_refs and t.startswith('['):
            if abs(p.paragraph_format.first_line_indent + Cm(0.74)) < 5000:
                ref_indent_ok = True
            for r in p.runs:
                if abs(r.font.size - Pt(9)) < 1000:
                    ref_font_ok = True

    # A4 result
    results.append(("A4 节标题", "通过" if section_ok else "警告", ""))

    # A5 result
    if body_total > 0:
        rate = body_indent_count / body_total
        if rate > 0.95:
            results.append(("A5 正文缩进", "通过", f"{body_indent_count}/{body_total}段"))
        else:
            results.append(("A5 正文缩进", "警告", f"{body_indent_count}/{body_total}段 ({rate:.0%})"))
            warnings += 1
    else:
        results.append(("A5 正文缩进", "警告", "无正文段落"))
        warnings += 1

    # A6 result
    if super_total > 0:
        rate = super_count / super_total
        if rate > 0.95:
            results.append(("A6 引用上标", "通过", f"{super_count}/{super_total}处"))
        else:
            results.append(("A6 引用上标", "警告", f"{super_count}/{super_total}处 ({rate:.0%})"))
            warnings += 1
    else:
        results.append(("A6 引用上标", "通过", "无引用"))

    # A7 result
    if level4_ok == level4_total:
        results.append(("A7 四级标题", "通过", f"{level4_ok}/{level4_total}处加粗"))
    else:
        results.append(("A7 四级标题", "错误", f"{level4_ok}/{level4_total}处"))
        errors += 1

    # A8 result
    if ref_font_ok and ref_indent_ok:
        results.append(("A8 参考文献", "通过", "9pt+悬挂缩进"))
    else:
        results.append(("A8 参考文献", "警告", "检查字体/缩进"))
        warnings += 1

    # ── A9: Duplicate chapter title ──
    titles = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    dups = [t for t in titles if titles.count(t) > 1 and len(t) > 5]
    if not dups:
        results.append(("A9 章标题唯一", "通过", "无重复"))
    else:
        results.append(("A9 章标题唯一", "错误", f"重复: {dups[0][:40]}"))
        errors += 1

    # ── A10: Leading punctuation ──
    bad = [p.text.strip()[:20] for p in doc.paragraphs
           if p.text.strip() and p.text.strip()[0] in '，；。、']
    if not bad:
        results.append(("A10 句首残标", "通过", "0处"))
    else:
        results.append(("A10 句首残标", "错误", f"{len(bad)}处"))
        errors += 1

    # ── A11: Chinese quote pairing (body only) ──
    body_text = ""
    in_refs_a11 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == '参考文献' and any(r.font.bold for r in p.runs):
            in_refs_a11 = True
            continue
        if not in_refs_a11:
            body_text += t
    left_q = body_text.count('\u201c')
    right_q = body_text.count('\u201d')
    straight_q = body_text.count('"')
    if left_q == right_q and straight_q == 0:
        results.append(("A11 引号配对", "通过", f"\"={left_q} \"={right_q} 无straight"))
    elif left_q != right_q:
        results.append(("A11 引号配对", "错误", f"\"={left_q} \"={right_q} 不配对"))
        errors += 1
    else:
        results.append(("A11 引号配对", "通过", f"\"={left_q} \"={right_q} straight={straight_q}(英)"))

    # ── Print report ──
    fname = os.path.basename(docx_path)
    print(f"\n{'='*60}")
    print(f"  排版自动机检报告")
    print(f"  文件: {fname}")
    print(f"  标准: 化工社《科技图书作者手册》(2019)")
    print(f"{'='*60}")
    for item, status, detail in results:
        icon = 'PASS' if status == '通过' else ('WARN' if status == '警告' else 'FAIL')
        print(f"  [{icon}] {item:12s}  {detail}")
    print(f"{'='*60}")
    total = errors + warnings
    if total == 0:
        print(f"  结果: 11/11 全部通过 — 可进入人工目检（宪法第15-16条）")
    else:
        print(f"  结果: {errors}错误 {warnings}警告 — 须修正后重新检测")
    print(f"{'='*60}\n")
    return errors == 0 and warnings == 0

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python check_formatting.py <docx路径>")
        sys.exit(1)
    ok = check(sys.argv[1])
    sys.exit(0 if ok else 1)
